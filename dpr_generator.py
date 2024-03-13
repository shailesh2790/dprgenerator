import streamlit as st
from docx import Document
from datetime import datetime

# Cache the function that creates the word document to prevent it from
# rerunning unnecessarily.
@st.cache(allow_output_mutation=True, show_spinner=False)
def create_word_document(data):
    doc = Document()
    doc.add_heading('Artificial Lift DPR', level=1)

    doc.add_heading('Dated:', level=3)
    doc.add_paragraph(data['date'])

    doc.add_heading('Attended wells on plunger lift system:', level=3)
    doc.add_paragraph(data['attended_wells'])

    doc.add_heading('Pressure Parameters', level=3)
    doc.add_paragraph(f"FCHP: {data['FCHP']}, FTHP: {data['FTHP']}")
    doc.add_paragraph(f"SCHP: {data['SCHP']}, Regulator pressure: {data['regulator_pressure']}")

    doc.add_heading('Operational / Plunger Parameters', level=3)
    doc.add_paragraph(f"Plunger Runtime: {data['plunger_runtime']}")
    doc.add_paragraph(f"Plunger Count: {data['plunger_count']}")
    doc.add_paragraph(f"Shut in Time: {data['shut_in_time']}")
    doc.add_paragraph(f"Open Time: {data['open_time']}")
    doc.add_paragraph(f"After Flow Time: {data['after_flow_time']}")
    doc.add_paragraph(f"Battery Voltage: {data['battery_voltage']}")
    doc.add_paragraph(f"Condition of Plunger Surface Assembly: {data['condition_of_plunger']}")


    # ... (The rest of your document creation logic)

    # Save the document in-memory
    filename = f"Artificial_Lift_DPR_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    doc.save(filename)
    return filename

# Main app
def main():
    st.title('Artificial Lift DPR Generator')

    # Use session state to store user input to prevent loss on rerun
    if 'data' not in st.session_state:
        st.session_state['data'] = {
            'date': None,
            'attended_wells': '',
            'FCHP': '',
            'FTHP': '',
            'SCHP': '',
            'regulator_pressure': '',
            'plunger_runtime': '',
            'plunger_count': '',
            'shut_in_time': '',
            'open_time': '',
            'after_flow_time': '',
            'battery_voltage': '',
            'condition_of_plunger': ''
        }

    # Use form for input to batch process on submit
    with st.form(key='dpr_form'):
      
        st.session_state['data']['attended_wells'] = st.text_input('Attended Wells', st.session_state['data']['attended_wells'])
        st.session_state['data']['FCHP'] = st.text_input('FCHP', st.session_state['data']['FCHP'])
        st.session_state['data']['FTHP'] = st.text_input('FTHP', st.session_state['data']['FTHP'])
        st.session_state['data']['SCHP'] = st.text_input('SCHP', st.session_state['data']['SCHP'])
        st.session_state['data']['regulator_pressure'] = st.text_input('regulator_pressure', st.session_state['data']['regulator_pressure'])
        st.session_state['data']['plunger_runtime'] = st.text_input('plunger_runtime', st.session_state['data']['plunger_runtime'])
        st.session_state['data']['plunger_count'] = st.text_input('plunger_count', st.session_state['data']['plunger_count'])
        st.session_state['data']['shut_in_time'] = st.text_input('shut_in_time', st.session_state['data']['shut_in_time'])
        st.session_state['data']['open_time'] = st.text_input('open_time', st.session_state['data']['open_time'])
        st.session_state['data']['after_flow_time'] = st.text_input('after_flow_time', st.session_state['data']['after_flow_time'])
        st.session_state['data']['battery_voltage'] = st.text_input('battery_voltage', st.session_state['data']['battery_voltage'])
        st.session_state['data']['condition_of_plunger'] = st.text_input('condition_of_plunger', st.session_state['data']['condition_of_plunger'])
        
        # ... (The rest of your input fields)

        submitted = st.form_submit_button('Generate Report')
        if submitted:
            # Generate document only after form submission
            report_path = create_word_document(st.session_state['data'])
            st.session_state['report_path'] = report_path

    # Provide a download link for the generated document
    if 'report_path' in st.session_state and st.session_state['report_path']:
        with open(st.session_state['report_path'], "rb") as file:
            st.download_button(
                label="Download Report",
                data=file,
                file_name=st.session_state['report_path'],
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == '__main__':
    main()

